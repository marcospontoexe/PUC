# -*- coding: utf-8 -*-
"""
Gerador do relatorio da atividade somativa - disciplina "IA na Gestao de Negocios" (PUCPR).

Produz:
    fig1_layout_cd.png      - layout do CD modelado como grafo de roteirizacao
    fig2_framework.png      - framework integrado das duas abordagens propostas
    relatorio_ia_picking.docx - relatorio final formatado segundo a ABNT

O texto do relatorio vive em conteudo.py; este arquivo cuida apenas das figuras
e da formatacao. Para revisar o texto, edite conteudo.py e execute novamente:

    python gerar_relatorio.py
"""

import os
import re

import matplotlib
matplotlib.use("Agg")  # backend sem interface grafica (nao abre janela)
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyBboxPatch

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

import conteudo as C

# Diretorio deste script - garante que os arquivos sejam gravados ao lado dele
BASE = os.path.dirname(os.path.abspath(__file__))

FONTE = "Arial"          # fonte padrao do documento (ABNT aceita Arial ou Times)
CORPO = Pt(12)           # tamanho do corpo de texto
MENOR = Pt(10)           # legendas de quadros e figuras
MINIMO = Pt(9)           # conteudo das celulas dos quadros


# ===========================================================================
# 1. FIGURAS
# ===========================================================================

def gerar_figura_1(caminho):
    """Desenha o CD como grafo: estruturas porta-paletes, corredores, doca e
    uma rota de separacao em S-shape passando pelos enderecos de um lote."""
    fig, ax = plt.subplots(figsize=(9.2, 4.8))

    # Linhas de centro dos seis corredores representados
    corredores = [1.2, 2.8, 4.4, 6.0, 7.6, 9.2]

    # Estruturas porta-paletes: dois blocos flanqueando cada corredor
    for xc in corredores:
        for deslocamento in (-0.62, 0.14):
            ax.add_patch(Rectangle((xc + deslocamento, 1.0), 0.48, 3.4,
                                   facecolor="#dcdcdc", edgecolor="#8c8c8c",
                                   linewidth=0.8, zorder=1))

    # Corredores transversais (inferior e superior), por onde se muda de rua
    for y in (0.8, 4.6):
        ax.plot([0.3, 10.1], [y, y], linestyle="--", color="#9e9e9e",
                linewidth=1.0, zorder=1)

    # Doca de expedicao: ponto de partida e chegada de toda viagem
    ax.add_patch(Rectangle((0.05, 0.55), 0.9, 0.5, facecolor="#2e7d32",
                           edgecolor="#1b5e20", linewidth=1.0, zorder=3))
    ax.text(0.5, 0.8, "DOCA", ha="center", va="center", color="white",
            fontsize=7.5, fontweight="bold", zorder=4)

    # Rota em S-shape: sobe a rua 1, desce a 2, sobe a 3, desce a 4 e retorna
    rota = [(0.5, 0.8), (1.2, 0.8), (1.2, 4.6), (2.8, 4.6), (2.8, 0.8),
            (4.4, 0.8), (4.4, 4.6), (6.0, 4.6), (6.0, 0.8), (0.5, 0.8)]
    rx, ry = zip(*rota)
    ax.plot(rx, ry, color="#1565c0", linewidth=2.0, zorder=2,
            solid_capstyle="round")

    # Enderecos de picking que compoem o lote a ser separado
    enderecos = [(1.2, 2.2), (1.2, 3.6), (2.8, 3.1), (2.8, 1.5),
                 (4.4, 2.4), (4.4, 4.0), (6.0, 3.3), (6.0, 1.8)]
    ex, ey = zip(*enderecos)
    ax.scatter(ex, ey, s=52, color="#c62828", edgecolor="white",
               linewidth=1.0, zorder=5)

    # Identificacao das ruas
    for i, xc in enumerate(corredores, start=1):
        ax.text(xc, 0.28, f"Rua {i}", ha="center", va="center", fontsize=7.5,
                color="#424242")
    ax.text(10.05, 0.28, "...", ha="center", va="center", fontsize=9,
            color="#424242")

    # Legenda montada manualmente (os elementos vem de tipos graficos distintos)
    itens = [
        plt.Line2D([], [], marker="s", linestyle="", markersize=9,
                   markerfacecolor="#dcdcdc", markeredgecolor="#8c8c8c",
                   label="Estrutura porta-paletes"),
        plt.Line2D([], [], marker="o", linestyle="", markersize=7,
                   markerfacecolor="#c62828", markeredgecolor="white",
                   label="Endereços de picking do lote"),
        plt.Line2D([], [], color="#1565c0", linewidth=2.0,
                   label="Rota de separação"),
        plt.Line2D([], [], color="#9e9e9e", linestyle="--", linewidth=1.0,
                   label="Corredor transversal"),
        plt.Line2D([], [], marker="s", linestyle="", markersize=9,
                   markerfacecolor="#2e7d32", markeredgecolor="#1b5e20",
                   label="Doca de expedição"),
    ]
    ax.legend(handles=itens, loc="upper center", bbox_to_anchor=(0.5, -0.02),
              ncol=3, frameon=False, fontsize=8)

    ax.set_xlim(0, 10.4)
    ax.set_ylim(-0.55, 5.0)
    ax.axis("off")
    fig.tight_layout()
    fig.savefig(caminho, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def _caixa(ax, x, y, largura, altura, texto, cor_fundo, cor_borda, tamanho=7.6):
    """Desenha uma caixa arredondada com texto centralizado (usada na Figura 2)."""
    ax.add_patch(FancyBboxPatch((x, y), largura, altura,
                                boxstyle="round,pad=0.02,rounding_size=0.12",
                                facecolor=cor_fundo, edgecolor=cor_borda,
                                linewidth=1.1, zorder=2))
    ax.text(x + largura / 2, y + altura / 2, texto, ha="center", va="center",
            fontsize=tamanho, color="#212121", zorder=3, linespacing=1.35)


def gerar_figura_2(caminho):
    """Desenha o framework integrado em tres faixas: ciclo tatico (Abordagem B),
    ciclo operacional (Abordagem A) e a base de dados construida pelo KDD."""
    fig, ax = plt.subplots(figsize=(11.0, 6.2))

    az_fundo, az_borda = "#e3f0fb", "#1565c0"   # ciclo tatico
    am_fundo, am_borda = "#fdf0dc", "#e07b00"   # ciclo operacional
    ci_fundo, ci_borda = "#eeeeee", "#757575"   # base de dados

    # ---------------- Faixa superior: ciclo tatico (Abordagem B) -------------
    ax.text(1.3, 6.62, "CICLO TÁTICO  ·  Abordagem B  ·  periodicidade mensal/sazonal",
            fontsize=9, fontweight="bold", color=az_borda, ha="left", va="center")

    etapas_b = [
        "1. Classificação\ndo padrão de demanda\n(ADI / CV²)",
        "2. Previsão por SKU\nHolt-Winters · ARIMA\nCroston · SBA",
        "3. Curva ABC dinâmica\nk-médias · FP-Growth\n(afinidade entre SKUs)",
        "4. Realocação de\nendereços\n(programação inteira)",
        "5. Predição do tempo\nde tarefa\n(Random Forest)",
    ]
    for i, texto in enumerate(etapas_b):
        x = 1.3 + i * 2.15
        _caixa(ax, x, 5.25, 1.9, 1.15, texto, az_fundo, az_borda)
        if i < len(etapas_b) - 1:  # seta ligando a etapa seguinte
            ax.annotate("", xy=(x + 2.13, 5.82), xytext=(x + 1.92, 5.82),
                        arrowprops=dict(arrowstyle="->", color=az_borda, lw=1.2))

    # ---------------- Faixa central: ciclo operacional (Abordagem A) ---------
    ax.text(1.3, 3.78, "CICLO OPERACIONAL  ·  Abordagem A  ·  a cada onda de separação",
            fontsize=9, fontweight="bold", color=am_borda, ha="left", va="center")

    etapas_a = [
        "1. Grafo do CD +\nmatriz de distâncias\n(Floyd–Warshall)",
        "2. Formação de lotes\n(algoritmo genético)",
        "3. Roteirização\n(colônia de formigas)",
        "4. Camada de\ncongestionamento\ne reotimização on-line",
    ]
    for i, texto in enumerate(etapas_a):
        x = 1.3 + i * 2.73
        _caixa(ax, x, 2.35, 2.3, 1.15, texto, am_fundo, am_borda)
        if i < len(etapas_a) - 1:
            ax.annotate("", xy=(x + 2.71, 2.92), xytext=(x + 2.32, 2.92),
                        arrowprops=dict(arrowstyle="->", color=am_borda, lw=1.2))

    # ---------------- Faixa inferior: base de dados / KDD -------------------
    _caixa(ax, 1.3, 0.60, 10.5, 0.90,
           "BASE DE DADOS INTEGRADA  ·  WMS + ERP + coletores RF/RFID  →  Data Warehouse\n"
           "Processo KDD: seleção · pré-processamento · transformação",
           ci_fundo, ci_borda, tamanho=8.2)

    # Alimentacao da base para o ciclo operacional
    for x in (4.0, 9.0):
        ax.annotate("", xy=(x, 2.33), xytext=(x, 1.52),
                    arrowprops=dict(arrowstyle="->", color=ci_borda, lw=1.2))

    # Alimentacao da base para o ciclo tatico, contornando pela margem esquerda
    ax.annotate("", xy=(1.28, 5.82), xytext=(0.75, 1.05),
                arrowprops=dict(arrowstyle="->", color=ci_borda, lw=1.2,
                                connectionstyle="angle,angleA=90,angleB=0,rad=6"))

    # Acoplamento entre os dois ciclos: o mapa de enderecos desce, o realizado sobe.
    # A seta para baixo para acima do titulo da faixa operacional, para nao cruza-lo.
    ax.annotate("", xy=(3.2, 4.02), xytext=(3.2, 5.23),
                arrowprops=dict(arrowstyle="->", color="#455a64", lw=1.6))
    ax.text(3.35, 4.60, "mapa de endereçamento\natualizado", fontsize=7.6,
            color="#455a64", ha="left", va="center", linespacing=1.3)

    ax.annotate("", xy=(9.9, 5.23), xytext=(9.9, 3.55),
                arrowprops=dict(arrowstyle="->", color="#455a64", lw=1.6,
                                linestyle="dashed"))
    ax.text(9.75, 4.42, "demanda, tempos e\ndistâncias realizados", fontsize=7.6,
            color="#455a64", ha="right", va="center", linespacing=1.3)

    ax.set_xlim(0.3, 12.0)
    ax.set_ylim(0.2, 7.0)
    ax.axis("off")
    fig.tight_layout()
    fig.savefig(caminho, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# ===========================================================================
# 2. FORMATACAO ABNT
# ===========================================================================

# Reconhece **negrito** e _italico_ dentro do texto dos paragrafos
MARCACAO = re.compile(r"(\*\*.+?\*\*|_[^_]+?_)")


def _aplicar_runs(paragrafo, texto, tamanho=CORPO):
    """Quebra o texto na marcacao inline e cria um run por trecho, aplicando
    negrito ou italico conforme o delimitador encontrado."""
    for trecho in MARCACAO.split(texto):
        if not trecho:
            continue
        if trecho.startswith("**") and trecho.endswith("**"):
            run = paragrafo.add_run(trecho[2:-2])
            run.bold = True
        elif trecho.startswith("_") and trecho.endswith("_"):
            run = paragrafo.add_run(trecho[1:-1])
            run.italic = True
        else:
            run = paragrafo.add_run(trecho)
        run.font.name = FONTE
        run.font.size = tamanho


def configurar_documento(doc):
    """Aplica as margens e o estilo padrao exigidos pela ABNT."""
    secao = doc.sections[0]
    secao.top_margin = Cm(3)
    secao.left_margin = Cm(3)
    secao.bottom_margin = Cm(2)
    secao.right_margin = Cm(2)

    estilo = doc.styles["Normal"]
    estilo.font.name = FONTE
    estilo.font.size = CORPO
    # Garante a fonte tambem para caracteres tratados como "east asian" pelo Word
    estilo.element.rPr.rFonts.set(qn("w:eastAsia"), FONTE)

    formato = estilo.paragraph_format
    formato.line_spacing = 1.5
    formato.space_before = Pt(0)
    formato.space_after = Pt(0)
    formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_paragrafo(doc, texto):
    """Paragrafo de corpo: justificado, entrelinhas 1,5 e recuo de 1,25 cm."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    _aplicar_runs(p, texto)
    return p


def add_citacao(doc, texto):
    """Citacao longa ABNT: recuo de 4 cm, fonte 10 e entrelinhas simples."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Cm(4)
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(12)
    _aplicar_runs(p, texto, tamanho=MENOR)
    return p


def add_titulo(doc, texto, nivel):
    """Titulo de secao numerada. Nivel 1 em caixa alta e negrito, nivel 2 em
    negrito e nivel 3 sem destaque - todos alinhados a esquerda, sem recuo."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(18 if nivel == 1 else 12)
    fmt.space_after = Pt(6)
    fmt.keep_with_next = True

    run = p.add_run(texto)
    run.font.name = FONTE
    run.font.size = CORPO
    run.bold = nivel in (1, 2)
    return p


def _legenda(doc, texto, tamanho=MENOR, centralizar=False, espaco_antes=12,
             espaco_depois=3, negrito_prefixo=None):
    """Paragrafo auxiliar usado nos titulos e nas fontes de quadros e figuras."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment = (WD_ALIGN_PARAGRAPH.CENTER if centralizar
                     else WD_ALIGN_PARAGRAPH.LEFT)
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(espaco_antes)
    fmt.space_after = Pt(espaco_depois)
    fmt.keep_with_next = espaco_depois <= 3

    if negrito_prefixo:
        run = p.add_run(negrito_prefixo)
        run.font.name = FONTE
        run.font.size = tamanho
        run.bold = True
    run = p.add_run(texto)
    run.font.name = FONTE
    run.font.size = tamanho
    return p


def _preencher_celula(celula, texto, negrito=False, centralizar=False):
    """Escreve o conteudo de uma celula sem interpretar marcacao inline (os
    quadros contem nomes de campos com underline, que nao devem virar italico)."""
    celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = celula.paragraphs[0]
    fmt = p.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(2)
    fmt.first_line_indent = Cm(0)
    fmt.alignment = (WD_ALIGN_PARAGRAPH.CENTER if centralizar
                     else WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(texto)
    run.font.name = FONTE
    run.font.size = MINIMO
    run.bold = negrito


def add_quadro(doc, dados):
    """Quadro ABNT: titulo acima, tabela com grade e fonte abaixo."""
    _legenda(doc, dados["titulo"], centralizar=False,
             negrito_prefixo=f"Quadro {dados['numero']} – ")

    cabecalho = dados["cabecalho"]
    linhas = dados["linhas"]
    larguras = dados["larguras"]

    tabela = doc.add_table(rows=len(linhas) + 1, cols=len(cabecalho))
    tabela.style = "Table Grid"
    tabela.autofit = False

    # Linha de cabecalho, em negrito e centralizada
    for j, titulo_col in enumerate(cabecalho):
        _preencher_celula(tabela.rows[0].cells[j], titulo_col,
                          negrito=True, centralizar=True)

    # Demais linhas: primeira coluna a esquerda, colunas numericas centralizadas
    for i, linha in enumerate(linhas, start=1):
        for j, valor in enumerate(linha):
            _preencher_celula(tabela.rows[i].cells[j], valor,
                              centralizar=(j > 0 and len(cabecalho) >= 4))

    # Larguras fixas precisam ser aplicadas celula a celula no python-docx
    for linha in tabela.rows:
        for j, celula in enumerate(linha.cells):
            celula.width = Cm(larguras[j])

    _legenda(doc, C.FONTE_AUTOR, tamanho=MINIMO, espaco_antes=3,
             espaco_depois=12)


def add_figura(doc, dados):
    """Figura ABNT: legenda acima, imagem centralizada e fonte abaixo."""
    _legenda(doc, dados["titulo"], centralizar=True,
             negrito_prefixo=f"Figura {dados['numero']} – ")

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(os.path.join(BASE, dados["arquivo"]),
                            width=Cm(dados["largura_cm"]))

    _legenda(doc, C.FONTE_AUTOR, tamanho=MINIMO, centralizar=True,
             espaco_antes=0, espaco_depois=12)


def montar_capa(doc):
    """Capa simples no padrao ABNT, encerrada por quebra de pagina."""
    def centrado(texto, negrito=False, tamanho=CORPO, antes=0, depois=0):
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = Cm(0)
        fmt.space_before = Pt(antes)
        fmt.space_after = Pt(depois)
        run = p.add_run(texto)
        run.font.name = FONTE
        run.font.size = tamanho
        run.bold = negrito
        return p

    centrado(C.INSTITUICAO, negrito=True)
    centrado(C.CURSO)
    centrado(C.DISCIPLINA, depois=180)
    centrado(C.TITULO, negrito=True, depois=180)
    centrado(C.AUTOR, depois=180)
    for parte in C.LOCAL_ANO.split("\n"):
        centrado(parte)

    quebra = doc.add_paragraph()
    quebra.add_run().add_break(WD_BREAK.PAGE)


def montar_referencias(doc):
    """Secao de referencias: alfabetica, alinhada a esquerda, entrelinhas
    simples e separada por espaco entre as entradas."""
    add_titulo(doc, "REFERÊNCIAS", nivel=1)
    for item in C.REFERENCIAS:
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Cm(0)
        fmt.line_spacing = 1.0
        fmt.space_after = Pt(12)
        _aplicar_runs(p, item)


# ===========================================================================
# 3. MONTAGEM DO DOCUMENTO
# ===========================================================================

def main():
    fig1 = os.path.join(BASE, "fig1_layout_cd.png")
    fig2 = os.path.join(BASE, "fig2_framework.png")
    gerar_figura_1(fig1)
    gerar_figura_2(fig2)
    print("Figuras geradas.")

    doc = Document()
    configurar_documento(doc)
    montar_capa(doc)

    # Percorre os blocos de conteudo e delega ao helper correspondente
    despacho = {
        "t1": lambda d: add_titulo(doc, d, 1),
        "t2": lambda d: add_titulo(doc, d, 2),
        "t3": lambda d: add_titulo(doc, d, 3),
        "p": lambda d: add_paragrafo(doc, d),
        "cit": lambda d: add_citacao(doc, d),
        "quadro": lambda d: add_quadro(doc, d),
        "figura": lambda d: add_figura(doc, d),
    }
    for tipo, dado in C.BLOCOS:
        despacho[tipo](dado)

    montar_referencias(doc)

    saida = os.path.join(BASE, "relatorio_ia_picking.docx")
    doc.save(saida)
    print("Relatorio gravado em:", saida)


if __name__ == "__main__":
    main()
